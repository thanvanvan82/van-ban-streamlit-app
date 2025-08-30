# app.py

import os
import re
import io

import dash
from dash import dcc, html, callback, Input, Output, State, ALL, ctx
import dash_bootstrap_components as dbc
from docxtpl import DocxTemplate
from docx import Document

# ==================== CẤU HÌNH VÀ HẰNG SỐ (ĐÃ CẬP NHẬT) ====================

DOCUMENT_TYPES = {
    "Văn bản hành chính TLU": {
        "Mẫu 1: Nghị quyết (cá biệt)": {"list": "tlu/list_1.docx", "template": "tlu/mau_1_nghi_quyet.docx"},
        "Mẫu 2: Quyết định (Hội đồng, Ban)": {"list": "tlu/list_2.docx", "template": "tlu/mau_2_qd_truc_tiep_hd_ban.docx"},
        "Mẫu 3: Quyết định (trực tiếp)": {"list": "tlu/list_3.docx", "template": "tlu/mau_3_qd_truc_tiep.docx"},
        "Mẫu 4: Quyết định (gián tiếp)": {"list": "tlu/list_4.docx", "template": "tlu/mau_4_qd_ban_hanh.docx"},
        "Mẫu 4a: Quy chế, quy định": {"list": "tlu/list_4a.docx", "template": "tlu/mau_4a_quy_che.docx"},
        "Mẫu 4b: Văn bản khác (kèm theo)": {"list": "tlu/list_4b.docx", "template": "tlu/mau_4b_vb_khac.docx"},
        "Mẫu 5a: Công văn (gửi 1 đơn vị)": {"list": "tlu/list_5a.docx", "template": "tlu/mau_5a_cong_van_hanh_chinh_mot.docx"},
        "Mẫu 5b: Công văn (gửi nhiều đơn vị)": {"list": "tlu/list_5b.docx", "template": "tlu/mau_5b_cong_van_hanh_chinh_nhieu.docx"},
        "Mẫu 6: Văn bản có tên loại": {"list": "tlu/list_6.docx", "template": "tlu/mau_6_thong_bao_bao_cao.docx"},
        "Mẫu 7: Tờ trình": {"list": "tlu/list_7.docx", "template": "tlu/mau_7_to_trinh.docx"},
        "Mẫu 8: Biên bản": {"list": "tlu/list_8.docx", "template": "tlu/mau_8_bien_ban.docx"},
        "Mẫu 9a: Giấy mời": {"list": "tlu/list_9a.docx", "template": "tlu/mau_9a_giay_moi.docx"},
        "Mẫu 9b: Giấy mời": {"list": "tlu/list_9b.docx", "template": "tlu/mau_9b_giay_moi.docx"},
        "Mẫu 10: Phụ lục": {"list": "tlu/list_10.docx", "template": "tlu/mau_10_phu_luc.docx"},
    },
    "Văn bản Đào tạo": {
        "ĐT-01: Đề cương chi tiết học phần": {"list": "dao_tao/list_de_cuong.docx", "template": "dao_tao/mau_de_cuong.docx"},
        "ĐT-02: Lịch trình giảng dạy": {"list": "dao_tao/list_lich_trinh.docx", "template": "dao_tao/mau_lich_trinh.docx"},
        "ĐT-03: Biên bản họp bộ môn": {"list": "dao_tao/list_bien_ban_hop.docx", "template": "dao_tao/mau_bien_ban_hop.docx"},
    },
    "Văn bản Thi đua": {
        "TĐ-01: Đăng ký thi đua": {"list": "thi_dua/list_dang_ky.docx", "template": "thi_dua/mau_dang_ky.docx"},
        "TĐ-02: Phiếu đánh giá viên chức": {"list": "thi_dua/list_phieu_danh_gia.docx", "template": "thi_dua/mau_phieu_danh_gia.docx"},
        "TĐ-03: Thành tích thay thế sáng kiến": {"list": "thi_dua/list_sang_kien.docx", "template": "thi_dua/mau_sang_kien.docx"},
    },
    "Văn bản Đảng": {
        "Đ-01: Bản kiểm điểm cá nhân": {"list": "dang/list_kiem_diem.docx", "template": "dang/mau_kiem_diem.docx"},
    },
    "Văn bản Hội đồng": {
        "HĐ-01: Đồ án tốt nghiệp": {"list": "hoi_dong/list_do_an_tn.docx", "template": "hoi_dong/mau_do_an_tn.docx"},
        "HĐ-02: Đề cương đề án NCS": {"list": "hoi_dong/list_de_cuong_ncs.docx", "template": "hoi_dong/mau_de_cuong_ncs.docx"},
        "HĐ-03: Bản thảo đề án NCS": {"list": "hoi_dong/list_ban_thao_ncs.docx", "template": "hoi_dong/mau_ban_thao_ncs.docx"},
        "HĐ-04: Biên bản bảo vệ đề án NCS": {"list": "hoi_dong/list_bao_ve_ncs.docx", "template": "hoi_dong/mau_bao_ve_ncs.docx"},
        "HĐ-05: Biên bản xét tuyển NCS": {"list": "hoi_dong/list_xet_tuyen_ncs.docx", "template": "hoi_dong/mau_xet_tuyen_ncs.docx"},
        "HĐ-06: Tiểu ban đánh giá chuyên đề": {"list": "hoi_dong/list_danh_gia_cd.docx", "template": "hoi_dong/mau_danh_gia_cd.docx"},
    },
    "Văn bản Khác": {
        "VK-01: Báo cáo sinh hoạt học thuật": {"list": "khac/list_shht.docx", "template": "khac/mau_shht.docx"},
    }
}

DEFAULT_GROUP = "Văn bản hành chính TLU"
DEFAULT_TEMPLATE = "Mẫu 5a: Công văn (gửi 1 đơn vị)"
DATA_DIR = "data"

app = dash.Dash(__name__, external_stylesheets=[dbc.themes.BOOTSTRAP], suppress_callback_exceptions=True)
app.title = "Tạo văn bản tự động"
server = app.server

# ==================== CÁC HÀM HỖ TRỢ (Giữ nguyên) ====================
def read_fields_and_data_from_list_docx(file_path):
    try:
        doc = Document(file_path)
        fields, data = [], {}
        if doc.tables:
            for table in doc.tables:
                for row_idx, row in enumerate(table.rows):
                    cells = [cell.text.strip() for cell in row.cells]
                    if row_idx == 0 and any(h in cells[0].lower() for h in ['tên trường', 'field']): continue
                    if len(cells) >= 3 and cells[0] and cells[1]:
                        name, label, value = cells[0], cells[1], cells[2]
                        clean_name = re.sub(r'[^\w]', '_', name.lower()).strip('_')
                        if clean_name:
                            fields.append({'name': clean_name, 'label': label, 'type': 'text_area' if any(w in label.lower() for w in ['nội dung', 'trích yếu']) else 'text_input'})
                            if value: data[clean_name] = value
        if not fields:
            for p in doc.paragraphs:
                if ':' in p.text.strip():
                    key, value = map(str.strip, p.text.split(':', 1))
                    clean_key = re.sub(r'[^\w]', '_', key.lower()).strip('_')
                    if clean_key:
                        fields.append({'name': clean_key, 'label': key, 'type': 'text_input'})
                        data[clean_key] = value
        return fields, data
    except Exception: return [], {}

def extract_template_variables(template_path):
    try:
        doc = Document(template_path)
        text = "\n".join(p.text for p in doc.paragraphs)
        for t in doc.tables: text += "\n".join(c.text for r in t.rows for c in r.cells)
        return list(set(re.findall(r'\{\{\s*(\w+)\s*\}\}', text)))
    except Exception: return []

def create_dash_input_field(field):
    label, name, f_type = field['label'], field['name'], field.get('type', 'text_input')
    input_id = {'type': 'form-input', 'field_name': name}
    if f_type == 'text_area':
        component = dbc.Textarea(id=input_id, style={"height": "100px"})
    else:
        component = dbc.Input(id=input_id, type="text")
    return dbc.Form([dbc.Label(label), component], className="mb-3")

# ==================== BỐ CỤC ỨNG DỤNG (ĐÃ CẬP NHẬT) ====================
app.layout = dbc.Container([
    dcc.Store(id='fields-store'), dcc.Store(id='data-store'),
    dcc.Store(id='template-vars-store'), dcc.Store(id='file-paths-store'),
    dcc.Download(id="download-docx"),

    html.H1("📄 Tạo văn bản tự động", className="my-4 text-center"),
    
    dbc.Row([
        dbc.Col(dbc.Card([
            dbc.CardHeader("Chọn Mẫu trình bày"),
            dbc.CardBody([
                dbc.Form([
                    dbc.Label("1. Chọn Nhóm văn bản:"),
                    dcc.Dropdown(
                        id='group-dropdown',
                        options=[{'label': group, 'value': group} for group in DOCUMENT_TYPES.keys()],
                        value=DEFAULT_GROUP,
                        clearable=False
                    ),
                ], className="mb-3"),
                dbc.Form([
                    dbc.Label("2. Chọn Mẫu văn bản cụ thể:"),
                    dcc.Dropdown(id='template-dropdown', clearable=False) # Options sẽ được cập nhật bởi callback
                ]),
                html.Hr(),
                html.Div(id='file-status-display')
            ])
        ]), width=4),
        dbc.Col([
            html.Div(id='main-content-display'),
            html.Div(id='notification-area', className="mt-3")
        ], width=8)
    ])
], fluid=True)

# ==================== CALLBACKS (ĐÃ CẬP NHẬT) ====================

# Callback 1: Cập nhật dropdown mẫu khi nhóm thay đổi
@app.callback(
    Output('template-dropdown', 'options'),
    Output('template-dropdown', 'value'),
    Input('group-dropdown', 'value')
)
def update_template_dropdown(selected_group):
    if not selected_group: return [], None
    templates = DOCUMENT_TYPES.get(selected_group, {})
    options = [{'label': name, 'value': name} for name in templates.keys()]
    
    # Set giá trị mặc định cho dropdown template
    default_value = None
    if selected_group == DEFAULT_GROUP and DEFAULT_TEMPLATE in templates:
        default_value = DEFAULT_TEMPLATE
    elif options:
        default_value = options[0]['value']
        
    return options, default_value

# Callback 2: Phân tích file khi MẪU VĂN BẢN thay đổi
@app.callback(
    Output('file-paths-store', 'data'),
    Output('fields-store', 'data'),
    Output('data-store', 'data'),
    Output('template-vars-store', 'data'),
    Output('file-status-display', 'children'),
    Input('template-dropdown', 'value'),
    State('group-dropdown', 'value')
)
def analyze_files(selected_template, selected_group):
    if not selected_template or not selected_group:
        return {}, [], {}, [], dbc.Alert("Vui lòng chọn mẫu.", color="warning")

    try:
        selected_files = DOCUMENT_TYPES[selected_group][selected_template]
        list_path = os.path.join(DATA_DIR, selected_files["list"])
        template_path = os.path.join(DATA_DIR, selected_files["template"])
    except KeyError:
        return {}, [], {}, [], dbc.Alert("Lỗi cấu hình: không tìm thấy thông tin file.", color="danger")
    
    paths = {'list_path': list_path, 'template_path': template_path}
    
    # Kiểm tra file và tạo thông báo
    alerts = []
    template_exists = os.path.exists(template_path)
    list_exists = os.path.exists(list_path)
    
    alerts.append(dbc.Alert(f"✅ Template: {template_path}", color="success") if template_exists else dbc.Alert(f"❌ Thiếu file: {template_path}", color="danger"))
    alerts.append(dbc.Alert(f"✅ Dữ liệu: {list_path}", color="success") if list_exists else dbc.Alert(f"❌ Thiếu file: {list_path}", color="danger"))
    
    if not template_exists or not list_exists:
        return paths, [], {}, [], alerts

    fields, data = read_fields_and_data_from_list_docx(list_path)
    template_vars = extract_template_variables(template_path)
    
    return paths, fields, data, template_vars, alerts

# Callback 3: Hiển thị form/nút bấm (Giữ nguyên)
@app.callback(
    Output('main-content-display', 'children'),
    Input('fields-store', 'data'),
    Input('data-store', 'data'),
    State('file-paths-store', 'data')
)
def render_main_content(fields, data, paths):
    if not fields or not paths or not os.path.exists(paths.get('list_path', '')):
        return dbc.Alert("Chọn một mẫu hợp lệ để bắt đầu.", color="info")

    fields_without_data = [f for f in fields if f['name'] not in data or not data[f['name']]]
    
    if not fields_without_data:
        card_content = [
            html.P("Nhấn nút bên dưới để tạo văn bản ngay."),
            dbc.Button("🔧 Tạo văn bản từ dữ liệu có sẵn", id="generate-button", color="primary", n_clicks=0, className="w-100")
        ]
        return dbc.Card([dbc.CardHeader("🎉 Tất cả dữ liệu đã có sẵn!"), dbc.CardBody(card_content)])
    else:
        form_inputs = [create_dash_input_field(field) for field in fields_without_data]
        card_content = [
            *form_inputs,
            dbc.Button("🔧 Tạo văn bản", id="generate-button", color="primary", n_clicks=0, className="w-100")
        ]
        return dbc.Card([dbc.CardHeader("✍️ Vui lòng nhập dữ liệu còn thiếu"), dbc.CardBody(card_content)])

# Callback 4: Tạo và tải văn bản (Giữ nguyên)
@app.callback(
    Output('download-docx', 'data'),
    Output('notification-area', 'children'),
    Input('generate-button', 'n_clicks'),
    State('file-paths-store', 'data'),
    State('template-vars-store', 'data'),
    State('data-store', 'data'),
    State({'type': 'form-input', 'field_name': ALL}, 'value'),
    State({'type': 'form-input', 'field_name': ALL}, 'id'),
    prevent_initial_call=True
)
def generate_document(n_clicks, paths, template_vars, existing_data, form_values, form_ids):
    if n_clicks == 0: return dash.no_update, dash.no_update
    template_path = paths.get('template_path')
    if not template_path or not os.path.exists(template_path):
        return dash.no_update, dbc.Alert("Lỗi: Không tìm thấy file template.", color="danger")

    final_context = existing_data.copy()
    form_data = {id['field_name']: val for id, val in zip(form_ids, form_values)}
    final_context.update(form_data)
    for var in template_vars: final_context.setdefault(var, f"[{var}]")

    try:
        doc = DocxTemplate(template_path)
        doc.render(final_context)
        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        
        file_name = f"vb_{final_context.get('so_ky_hieu', 'output').replace('/', '_')}.docx"
        download_data = dcc.send_bytes(file_stream.read(), file_name)
        notification = dbc.Alert("✅ Tạo văn bản thành công!", color="success", duration=4000)
        return download_data, notification
    except Exception as e:
        return dash.no_update, dbc.Alert(f"❌ Lỗi khi tạo văn bản: {e}", color="danger")

# ==================== CHẠY ỨNG DỤNG ====================
if __name__ == '__main__':
    app.run_server(debug=True)